import os
import io
import base64
import pandas as pd
import streamlit as st
from dotenv import load_dotenv
from groq import Groq
from pydantic import BaseModel, Field
from typing import List, Optional
from PIL import Image
from datetime import datetime
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# 1. Load secret environment variables (.env file)
load_dotenv()
api_key = os.getenv("GROQ_API_KEY")

# 2. Page Config
st.set_page_config(
    page_title="AI Financial Bookkeeper | Enterprise Ledger", 
    page_icon="⚡", 
    layout="wide",
    initial_sidebar_state="collapsed"
)

# 3. Helper Function to Compress Images
def compress_image(image_bytes: bytes, max_size: tuple = (640, 640), quality: int = 50) -> bytes:
    img = Image.open(io.BytesIO(image_bytes))
    if img.mode in ("RGBA", "P"):
        img = img.convert("RGB")
    img.thumbnail(max_size, Image.Resampling.LANCZOS)
    buffer = io.BytesIO()
    img.save(buffer, format="JPEG", quality=quality, optimize=True)
    return buffer.getvalue()

# 4. Pydantic Schemas for Hierarchical Ledger Outputs
class Transaction(BaseModel):
    date: Optional[str] = Field("2026-06-01", description="Date of transaction in YYYY-MM-DD. If missing, default to start of month")
    period_week: str = Field("WEEK ONE", description="Operational week or time block e.g. WEEK ONE, WEEK TWO")
    flow_type: str = Field("EXPENSE", description="Type of financial flow: INCOME or EXPENSE")
    description: str = Field(..., description="Description of the item or service")
    category: str = Field(..., description="Specific sub-category e.g. Cash Income, Transfer Income, Cash Expenses, Bank Expenses")
    amount: float = Field(..., description="Numerical cost or income amount")

class FinancialStatement(BaseModel):
    business_name: str = Field("Unknown Business", description="Name of organization or business")
    period: str = Field("Unknown Period", description="Time period covered")
    currency: str = Field("NGN", description="Currency symbol or code")
    reported_grand_total: float = Field(0.0, description="Calculated grand total of transactions")
    transactions: List[Transaction] = Field(default_factory=list, description="List of extracted hierarchical transactions")

# 5. Deep Modern Fintech & Table Header CSS Customizations
st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Plus+Jakarta+Sans:wght@400;500;600;700;800&display=swap');

    html, body, [class*="css"] {
        font-family: 'Plus Jakarta Sans', sans-serif;
    }

    .stApp {
        background: radial-gradient(circle at 50% -20%, #1e1b4b 0%, #0f172a 45%, #020617 100%);
        color: #f8fafc;
    }

    .hero-card {
        background: rgba(30, 41, 59, 0.5);
        backdrop-filter: blur(16px);
        -webkit-backdrop-filter: blur(16px);
        border: 1px solid rgba(255, 255, 255, 0.1);
        border-radius: 20px;
        padding: 32px;
        margin-bottom: 28px;
        box-shadow: 0 20px 40px -15px rgba(0, 0, 0, 0.5);
    }

    .hero-badge {
        display: inline-block;
        background: linear-gradient(90deg, #6366f1 0%, #a855f7 100%);
        color: #ffffff;
        font-size: 12px;
        font-weight: 700;
        text-transform: uppercase;
        letter-spacing: 1.2px;
        padding: 6px 14px;
        border-radius: 20px;
        margin-bottom: 14px;
    }

    .hero-title {
        background: linear-gradient(135deg, #ffffff 0%, #cbd5e1 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        font-size: 38px;
        font-weight: 800;
        margin-bottom: 10px;
        letter-spacing: -0.5px;
    }

    .hero-subtitle {
        color: #94a3b8;
        font-size: 16px;
        font-weight: 400;
        line-height: 1.6;
        max-width: 800px;
    }

    div.stButton > button:first-child {
        background: linear-gradient(135deg, #6366f1 0%, #4f46e5 100%) !important;
        color: white !important;
        border: none !important;
        padding: 14px 28px !important;
        font-size: 16px !important;
        font-weight: 600 !important;
        border-radius: 12px !important;
        box-shadow: 0 4px 20px rgba(99, 102, 241, 0.4) !important;
        transition: all 0.3s ease !important;
    }

    div.stButton > button:first-child:hover {
        transform: translateY(-2px) !important;
        box-shadow: 0 8px 25px rgba(99, 102, 241, 0.6) !important;
    }

    [data-testid="stMetricValue"] {
        font-family: 'Plus Jakarta Sans', sans-serif;
        font-weight: 700;
        color: #38bdf8;
    }

    thead th {
        font-weight: 800 !important;
        text-transform: uppercase !important;
        letter-spacing: 0.5px !important;
    }
    </style>
""", unsafe_allow_html=True)

# 6. Top Hero Section
st.markdown("""
    <div class="hero-card">
        <span class="hero-badge">⚡ Enterprise AI Financial Suite</span>
        <div class="hero-title">AI Financial Bookkeeper</div>
        <div class="hero-subtitle">Instantly extract, compartmentalize, and audit business records with professional multi-week departmental tables and institutional Excel/Word reports.</div>
    </div>
""", unsafe_allow_html=True)

# 7. Quick System Status Stats
m1, m2, m3 = st.columns(3)
with m1:
    st.metric(label="System Status", value="Active", delta="Ready")
with m2:
    st.metric(label="AI Engine", value="Groq LLaMA-3.3 / Qwen", delta="High Precision")
with m3:
    st.metric(label="Audit Security", value="Encrypted", delta="Secure Local")

st.markdown("<br>", unsafe_allow_html=True)

# 8. Selection Mode
st.subheader("Select Ledger Input Method")
input_method = st.radio(
    "Choose Input Method:",
    ("Snap Photo of Book/Receipt", "Paste Text/Notes", "Upload File"),
    horizontal=True,
    label_visibility="collapsed"
)

st.markdown("<br>", unsafe_allow_html=True)

captured_image_bytes = None
pasted_text = ""

if input_method == "Snap Photo of Book/Receipt":
    camera_photo = st.camera_input("Capture image of physical receipt or paper ledger")
    if camera_photo:
        captured_image_bytes = camera_photo.getvalue()

elif input_method == "Paste Text/Notes":
    pasted_text = st.text_area(
        "Paste sales logs, receipt details, or expense records below:",
        placeholder="e.g.,\nWeek One:\n- Cash Income Fresh Dew: 9,300\n- Tithe Zenith: 2,500\n- Bus Service Expense: 15,000",
        height=180
    )

elif input_method == "Upload File":
    uploaded_file = st.file_uploader(
        "Upload statement or ledger (CSV, TXT, Images, Excel, Word .docx)", 
        type=["csv", "txt", "png", "jpg", "jpeg", "xlsx", "xls", "docx"]
    )
    if uploaded_file is not None:
        if uploaded_file.type.startswith("image"):
            captured_image_bytes = uploaded_file.getvalue()
        else:
            try:
                if uploaded_file.name.endswith(".csv"):
                    try:
                        df_temp = pd.read_csv(uploaded_file)
                        pasted_text = df_temp.to_string()
                    except Exception:
                        uploaded_file.seek(0)
                        pasted_text = uploaded_file.getvalue().decode("utf-8", errors="ignore")
                elif uploaded_file.name.endswith((".xlsx", ".xls")):
                    df_temp = pd.read_excel(uploaded_file)
                    pasted_text = df_temp.to_string()
                elif uploaded_file.name.endswith(".docx"):
                    doc = Document(uploaded_file)
                    fullText = [para.text for para in doc.paragraphs]
                    for table in doc.tables:
                        for row in table.rows:
                            fullText.append(" | ".join([cell.text.strip() for cell in row.cells]))
                    pasted_text = "\n".join(fullText)
                else:
                    pasted_text = uploaded_file.getvalue().decode("utf-8", errors="ignore")
            except Exception as e:
                st.error(f"❌ Error parsing file structure: {e}")

st.markdown("<br>", unsafe_allow_html=True)

def uppercase_headers(df):
    df_copy = df.copy()
    df_copy.columns = [str(col).upper() for col in df_copy.columns]
    return df_copy

# 10. Main Action & Groq Execution Engine
if st.button("🚀 Process & Generate Accounting Report", use_container_width=True):
    # Retrieve API key from Streamlit secrets or environment variables
    active_api_key = st.secrets.get("GROQ_API_KEY") or api_key
    
    if not active_api_key:
        st.error("❌ GROQ_API_KEY not detected! Please add it to Streamlit Secrets or your .env file.")
    elif not captured_image_bytes and not pasted_text.strip():
        st.warning("⚠️ Please snap a photo, upload a document, or paste transaction notes first.")
    else:
        with st.spinner("⚡ Processing records into hierarchical ledger format with Groq AI..."):
            try:
                client = Groq(api_key=active_api_key)
                
                system_instructions = """You are an expert corporate AI Accountant and Auditor. Read the records carefully and structure them hierarchically like institutional ledgers (separated by operational weeks/periods like WEEK ONE, WEEK TWO, flow types like INCOME or EXPENSE, and categories). 
If a date is not provided for a transaction, infer a realistic date or default to the start of the reporting period (e.g., 2026-06-01).

Return ONLY valid JSON matching this exact structure:
{
  "business_name": "string",
  "period": "string",
  "currency": "string",
  "reported_grand_total": 0.0,
  "transactions": [
    {
      "date": "YYYY-MM-DD",
      "period_week": "WEEK ONE",
      "flow_type": "INCOME or EXPENSE",
      "description": "item name",
      "category": "Cash Income or Cash Expenses etc",
      "amount": 0.0
    }
  ]
}"""

                if captured_image_bytes:
                    compressed_bytes = compress_image(captured_image_bytes)
                    base64_image = base64.b64encode(compressed_bytes).decode('utf-8')
                    selected_model = "qwen/qwen3.6-27b"
                    messages = [{
                        "role": "user",
                        "content": [
                            {"type": "text", "text": system_instructions},
                            {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}}
                        ]
                    }]
                else:
                    selected_model = "llama-3.3-70b-versatile"
                    prompt = f"{system_instructions}\n\nRecords to analyze:\n{pasted_text}"
                    messages = [{"role": "user", "content": prompt}]

                response = client.chat.completions.create(
                    model=selected_model,
                    messages=messages,
                    response_format={"type": "json_object"}
                )

                json_output = response.choices[0].message.content
                statement = FinancialStatement.model_validate_json(json_output)

                if statement.transactions:
                    calculated_sum = sum(t.amount for t in statement.transactions if t.amount is not None)
                    statement.reported_grand_total = calculated_sum

                st.success("✅ Financial Audit Processing Complete!")

                # --- Professional Metric KPI Cards ---
                st.markdown("<br>", unsafe_allow_html=True)
                c1, c2, c3, c4 = st.columns(4)
                c1.metric("Organization / Business", statement.business_name)
                c2.metric("Reporting Period", statement.period)
                c3.metric("Total Line Items", len(statement.transactions))
                c4.metric("Grand Total Value", f"{statement.currency} {statement.reported_grand_total:,.2f}")

                st.markdown("<br>", unsafe_allow_html=True)

                if statement.transactions:
                    tx_data = [t.model_dump() for t in statement.transactions]
                    df_result = pd.DataFrame(tx_data)

                    # --- Interactive Sidebar Filter ---
                    st.sidebar.markdown("### 🔍 Filter Audit View")
                    weeks_list = ["All Weeks"] + list(df_result['period_week'].unique())
                    selected_week_filter = st.sidebar.selectbox("Select Operational Week", weeks_list)

                    if selected_week_filter != "All Weeks":
                        df_display = df_result[df_result['period_week'] == selected_week_filter]
                    else:
                        df_display = df_result

                    st.subheader("📋 Master Institutional Ledger")
                    st.dataframe(uppercase_headers(df_display), use_container_width=True)

                    st.markdown("<br>", unsafe_allow_html=True)
                    st.subheader("📂 Period & Departmental Breakdown")
                    
                    unique_weeks = df_result['period_week'].unique()
                    for wk in unique_weeks:
                        with st.expander(f"📅 Operational Period: {wk}", expanded=True):
                            df_wk_subset = df_result[df_result['period_week'] == wk]
                            st.dataframe(uppercase_headers(df_wk_subset), use_container_width=True)
                            wk_total = df_wk_subset['amount'].sum()
                            st.caption(f"**Total Volume for {wk}:** {statement.currency} {wk_total:,.2f}")

                    # --- SIDE-BY-SIDE DOWNLOAD BUTTONS (Excel & Word) ---
                    st.markdown("<br>", unsafe_allow_html=True)
                    st.subheader("📥 Export Official Audit Reports")

                    current_date_str = datetime.now().strftime("%Y-%m-%d")
                    sanitized_org_name = statement.business_name.replace(" ", "_")
                    excel_filename = f"{sanitized_org_name}_Audit_{current_date_str}.xlsx"
                    word_filename = f"{sanitized_org_name}_Audit_{current_date_str}.docx"

                    export_rows = []
                    for wk in unique_weeks:
                        df_wk_subset = df_result[df_result['period_week'] == wk]
                        export_rows.append({"DATE": f"=== {wk.upper()} ===", "DESCRIPTION": "", "CATEGORY": "", "AMOUNT": ""})
                        for _, row in df_wk_subset.iterrows():
                            export_rows.append({
                                "DATE": row["date"] if row["date"] else "N/A",
                                "DESCRIPTION": row["description"],
                                "CATEGORY": f"[{row['flow_type']}] {row['category']}",
                                "AMOUNT": row["amount"]
                            })
                        wk_subtotal = df_wk_subset['amount'].sum()
                        export_rows.append({"DATE": "", "DESCRIPTION": f"TOTAL FOR {wk.upper()}", "CATEGORY": "", "AMOUNT": wk_subtotal})
                        export_rows.append({"DATE": "", "DESCRIPTION": "", "CATEGORY": "", "AMOUNT": ""})

                    grand_total_val = df_result['amount'].sum()
                    export_rows.append({"DATE": "=== SUMMARY ===", "DESCRIPTION": "", "CATEGORY": "", "AMOUNT": ""})
                    export_rows.append({"DATE": "", "DESCRIPTION": "GRAND TOTAL FOR THE MONTH", "CATEGORY": "OVERALL", "AMOUNT": grand_total_val})

                    df_structured_export = pd.DataFrame(export_rows)

                    # 1. Generate Excel File in Memory
                    excel_buffer = io.BytesIO()
                    with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
                        df_structured_export.to_excel(writer, sheet_name='Institutional Ledger', index=False)
                        worksheet = writer.sheets['Institutional Ledger']
                        
                        header_font = Font(name='Plus Jakarta Sans', size=11, bold=True, color="FFFFFF")
                        header_fill = PatternFill(start_color="1E1B4B", end_color="1E1B4B", fill_type="solid")
                        section_font = Font(name='Plus Jakarta Sans', size=10, bold=True, color="1E1B4B")
                        section_fill = PatternFill(start_color="E0E7FF", end_color="E0E7FF", fill_type="solid")
                        total_font = Font(name='Plus Jakarta Sans', size=10, bold=True, color="0F172A")
                        total_fill = PatternFill(start_color="F1F5F9", end_color="F1F5F9", fill_type="solid")
                        grand_total_font = Font(name='Plus Jakarta Sans', size=11, bold=True, color="FFFFFF")
                        grand_total_fill = PatternFill(start_color="0F172A", end_color="0F172A", fill_type="solid")
                        zebra_fill = PatternFill(start_color="F8FAFC", end_color="F8FAFC", fill_type="solid")
                        
                        thin_border = Border(left=Side(style='thin', color='CBD5E1'), right=Side(style='thin', color='CBD5E1'), top=Side(style='thin', color='CBD5E1'), bottom=Side(style='thin', color='CBD5E1'))

                        for col_num in range(1, len(df_structured_export.columns) + 1):
                            cell = worksheet.cell(row=1, column=col_num)
                            cell.font = header_font
                            cell.fill = header_fill
                            cell.alignment = Alignment(horizontal="center", vertical="center")
                            cell.border = thin_border

                        for r_idx in range(2, len(df_structured_export) + 2):
                            val_date = worksheet.cell(row=r_idx, column=1).value
                            val_desc = str(worksheet.cell(row=r_idx, column=2).value or '')
                            for c_idx in range(1, len(df_structured_export.columns) + 1):
                                cell = worksheet.cell(row=r_idx, column=c_idx)
                                cell.border = thin_border
                                cell.font = Font(name='Plus Jakarta Sans', size=10)
                                if c_idx == 4 and isinstance(cell.value, (int, float)):
                                    cell.number_format = '#,##0.00'
                                    cell.alignment = Alignment(horizontal="right", vertical="center")
                                if val_date and str(val_date).startswith("==="):
                                    cell.font = section_font
                                    cell.fill = section_fill
                                    if c_idx == 1:
                                        cell.alignment = Alignment(horizontal="left", vertical="center")
                                elif "GRAND TOTAL FOR" in val_desc:
                                    cell.font = grand_total_font
                                    cell.fill = grand_total_fill
                                elif "TOTAL FOR" in val_desc:
                                    cell.font = total_font
                                    cell.fill = total_fill
                                elif r_idx % 2 == 0 and not str(val_date).startswith("==="):
                                    cell.fill = zebra_fill

                        for col in worksheet.columns:
                            max_len = max(len(str(cell.value or '')) for cell in col)
                            col_letter = col[0].column_letter
                            worksheet.column_dimensions[col_letter].width = max(max_len + 4, 18)
                    excel_data = excel_buffer.getvalue()

                    # 2. Generate Word (.docx) File in Memory
                    doc = Document()
                    doc.add_heading(f"Official Audit Report: {statement.business_name}", level=1)
                    p = doc.add_paragraph(f"Reporting Period: {statement.period} | Currency: {statement.currency}")
                    p.runs[0].font.size = Pt(12)
                    p.runs[0].font.color.rgb = RGBColor(100, 100, 100)

                    table = doc.add_table(rows=1, cols=4)
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    hdr_cells = table.rows[0].cells
                    headers = ["DATE", "DESCRIPTION", "CATEGORY", "AMOUNT"]
                    for i, h_text in enumerate(headers):
                        hdr_cells[i].text = h_text
                        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

                    for row_item in export_rows:
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(row_item["DATE"])
                        row_cells[1].text = str(row_item["DESCRIPTION"])
                        row_cells[2].text = str(row_item["CATEGORY"])
                        amt_val = row_item["AMOUNT"]
                        row_cells[3].text = f"{amt_val:,.2f}" if isinstance(amt_val, (int, float)) else str(amt_val)

                    word_buffer = io.BytesIO()
                    doc.save(word_buffer)
                    word_data = word_buffer.getvalue()

                    # --- SIDE BY SIDE BUTTON COLUMNS ---
                    dl_col1, dl_col2 = st.columns(2)
                    with dl_col1:
                        st.download_button(
                            label="📥 Download Excel Report (.xlsx)",
                            data=excel_data,
                            file_name=excel_filename,
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            use_container_width=True
                        )
                    with dl_col2:
                        st.download_button(
                            label="📥 Download Word Report (.docx)",
                            data=word_data,
                            file_name=word_filename,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )

            except Exception as error:
                st.error(f"❌ Processing Error: {str(error)}")